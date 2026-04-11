#!/usr/bin/env python3
"""
md_to_docx.py — minimal markdown → docx converter for communicative-interaction logs.

Usage:
    python3 md_to_docx.py input.md output.docx

Supports:
- # H1, ## H2, ### H3 headings
- **bold** inline
- *italic* inline (when not collision with bold)
- `code` inline
- - bullet lists
- 1. numbered lists
- Blockquotes (> prefix)
- Horizontal rules (---)
- Plain paragraphs
- Code blocks (``` fenced)
- Tables (| col | col |)

Does NOT support:
- Nested lists beyond one level
- Images
- Inline links (rendered as plain text)

This is intentionally minimal; use pandoc if you need full fidelity.
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


BOLD_RE = re.compile(r'\*\*([^*]+)\*\*')
ITALIC_RE = re.compile(r'(?<!\*)\*([^*]+)\*(?!\*)')
CODE_RE = re.compile(r'`([^`]+)`')


def add_formatted_runs(paragraph, text):
    """Add runs to a paragraph with basic inline formatting."""
    # First pass: find all bold/italic/code spans and tokenize
    # Simple approach: process bold, then italic, then code within remaining text
    i = 0
    while i < len(text):
        # Try bold
        m = BOLD_RE.match(text, i)
        if m:
            run = paragraph.add_run(m.group(1))
            run.bold = True
            i = m.end()
            continue
        # Try code
        m = CODE_RE.match(text, i)
        if m:
            run = paragraph.add_run(m.group(1))
            run.font.name = 'Courier New'
            i = m.end()
            continue
        # Try italic
        m = ITALIC_RE.match(text, i)
        if m:
            run = paragraph.add_run(m.group(1))
            run.italic = True
            i = m.end()
            continue
        # Plain char
        paragraph.add_run(text[i])
        i += 1


def convert(md_path: Path, docx_path: Path):
    md = md_path.read_text(encoding='utf-8')
    doc = Document()

    # Set default font to support Korean
    styles = doc.styles
    for style_name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3'):
        try:
            style = styles[style_name]
            style.font.name = 'Malgun Gothic'
            style.font.size = Pt(11) if style_name == 'Normal' else None
        except KeyError:
            pass

    lines = md.splitlines()
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # Parse cells
        parsed = [[c.strip() for c in row.strip().strip('|').split('|')] for row in table_rows]
        # Skip separator rows (like |---|---|)
        parsed = [r for r in parsed if not all(re.match(r'^:?-+:?$', c) for c in r)]
        if not parsed:
            in_table = False
            table_rows = []
            return
        ncols = max(len(r) for r in parsed)
        table = doc.add_table(rows=len(parsed), cols=ncols)
        table.style = 'Light Grid Accent 1'
        for row_idx, row in enumerate(parsed):
            for col_idx in range(ncols):
                cell_text = row[col_idx] if col_idx < len(row) else ''
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_runs(p, cell_text)
                if row_idx == 0:
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()  # spacer
        in_table = False
        table_rows = []

    for line in lines:
        # Code block fences
        if line.strip().startswith('```'):
            if in_code_block:
                # Close block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_block_lines = []
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # Table row
        if line.lstrip().startswith('|') and '|' in line.lstrip()[1:]:
            in_table = True
            table_rows.append(line)
            continue
        elif in_table:
            flush_table()

        # Horizontal rule
        if re.match(r'^\s*---+\s*$', line) or re.match(r'^\s*\*\*\*+\s*$', line):
            flush_table()
            p = doc.add_paragraph()
            run = p.add_run('─' * 50)
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            flush_table()
            level = min(len(m.group(1)), 3)
            heading = doc.add_heading(level=level)
            add_formatted_runs(heading, m.group(2))
            continue

        # Blockquote
        if line.lstrip().startswith('>'):
            flush_table()
            text = line.lstrip().lstrip('>').lstrip()
            p = doc.add_paragraph(style='Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else 'Normal')
            add_formatted_runs(p, text)
            continue

        # Bullet list
        m = re.match(r'^(\s*)[-*+]\s+(.*)$', line)
        if m:
            flush_table()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, m.group(2))
            continue

        # Numbered list
        m = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
        if m:
            flush_table()
            p = doc.add_paragraph(style='List Number')
            add_formatted_runs(p, m.group(2))
            continue

        # Empty line
        if not line.strip():
            continue

        # Plain paragraph
        flush_table()
        p = doc.add_paragraph()
        add_formatted_runs(p, line)

    flush_table()
    doc.save(str(docx_path))


def main():
    if len(sys.argv) != 3:
        print("Usage: md_to_docx.py input.md output.docx", file=sys.stderr)
        sys.exit(1)
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    if not md_path.exists():
        print(f"ERROR: input file not found: {md_path}", file=sys.stderr)
        sys.exit(1)
    convert(md_path, docx_path)
    print(f"OK: {md_path} → {docx_path}")


if __name__ == '__main__':
    main()
