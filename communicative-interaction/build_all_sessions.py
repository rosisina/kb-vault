#!/usr/bin/env python3
"""
build_all_sessions.py — merge every session markdown in this folder into ALL-SESSIONS.docx

Reads all files matching `*-session-*.md` in sorted order and concatenates them into
a single docx with a title page and per-session page breaks.

Usage:
    python3 build_all_sessions.py
"""

import sys
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_BREAK
except ImportError:
    print("ERROR: python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# Import the converter from the sibling script
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR))
from md_to_docx import add_formatted_runs


def main():
    folder = SCRIPT_DIR
    session_files = sorted(folder.glob('*-session-*.md'))
    if not session_files:
        print("No session files found.", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # Set default font to support Korean
    styles = doc.styles
    for style_name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3'):
        try:
            style = styles[style_name]
            style.font.name = 'Malgun Gothic'
            style.font.size = Pt(11) if style_name == 'Normal' else None
        except KeyError:
            pass

    # Title page
    title = doc.add_heading(level=0)
    title.add_run('Communicative Interaction — All Sessions')
    meta = doc.add_paragraph()
    meta.add_run(f'Cumulative log of James ↔ Claude interactions on knowledge-base / defense-kg-platform\n')
    meta.add_run(f'Generated: {date.today().isoformat()}\n')
    meta.add_run(f'Sessions included: {len(session_files)}\n')
    for idx, sf in enumerate(session_files, 1):
        meta.add_run(f'  {idx:02d}. {sf.stem}\n')
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Per-session content
    for idx, sf in enumerate(session_files):
        md = sf.read_text(encoding='utf-8')
        # Add a section divider heading
        divider = doc.add_heading(level=1)
        divider.add_run(f'[Session {idx + 1:02d}] {sf.stem}')

        # Simple inline conversion (abridged markdown)
        lines = md.splitlines()
        in_code = False
        code_lines = []
        for line in lines:
            if line.strip().startswith('```'):
                if in_code:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                continue
            if in_code:
                code_lines.append(line)
                continue

            import re
            m = re.match(r'^(#{1,6})\s+(.*)$', line)
            if m:
                lvl = min(len(m.group(1)), 3)
                h = doc.add_heading(level=lvl + 1 if lvl < 3 else 3)  # shift by one so top-level heading becomes level 2 under session heading
                add_formatted_runs(h, m.group(2))
                continue

            m = re.match(r'^(\s*)[-*+]\s+(.*)$', line)
            if m:
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_runs(p, m.group(2))
                continue

            m = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
            if m:
                p = doc.add_paragraph(style='List Number')
                add_formatted_runs(p, m.group(2))
                continue

            if line.lstrip().startswith('>'):
                text = line.lstrip().lstrip('>').lstrip()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                continue

            if not line.strip():
                continue

            p = doc.add_paragraph()
            add_formatted_runs(p, line)

        # Page break between sessions (except last)
        if idx < len(session_files) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    out = folder / 'ALL-SESSIONS.docx'
    doc.save(str(out))
    print(f'OK: {len(session_files)} sessions → {out}')


if __name__ == '__main__':
    main()
